from __future__ import annotations

from pathlib import Path
from typing import List, Optional
from decimal import Decimal
import logging

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scr.text_utils import format_amount


def _replace_text_in_paragraph(paragraph, placeholder: str, value: str) -> bool:
    # 通过 paragraph.text 统一替换
    text = paragraph.text or ""
    if placeholder in text:
        paragraph.text = text.replace(placeholder, value)
        return True
    return False


def _replace_text_in_table(table, placeholder: str, value: str) -> bool:
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if _replace_text_in_paragraph(p, placeholder, value):
                    replaced = True
    return replaced


def replace_placeholders(doc: Document, mapping: dict, logger: logging.Logger) -> None:
    for p in doc.paragraphs:
        for k, v in mapping.items():
            _replace_text_in_paragraph(p, k, v)
    for t in doc.tables:
        for k, v in mapping.items():
            _replace_text_in_table(t, k, v)


def _apply_table_style_or_borders(table) -> None:
    # 优先使用内置网格样式，失败则手动加边框
    try:
        table.style = 'Table Grid'
        return
    except Exception:
        pass
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.insert(0, OxmlElement('w:tblPr'))
    borders = OxmlElement('w:tblBorders')
    for tag in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')  # 0.5pt
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _center_table(table) -> None:
    # 优先 API 居中，失败则设置 XML jc=center
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return
    except Exception:
        pass
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.insert(0, OxmlElement('w:tblPr'))
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)


def _fill_table(table, headers: List[str], rows: List[List[str]]) -> None:
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = h
    for r in rows:
        tr = table.add_row().cells
        for idx, val in enumerate(r):
            tr[idx].text = val


def _apply_font_size_to_document(doc: Document, pt_size: int = 14, line_spacing_pt: int = 25) -> None:
    # 统一字号与行距（跳过“Title”样式段落）
    for p in doc.paragraphs:
        if p.style and str(p.style.name).lower() == 'title':
            continue
        for run in p.runs:
            run.font.size = Pt(pt_size)
        try:
            p.paragraph_format.line_spacing = Pt(line_spacing_pt)
        except Exception:
            pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.style and str(p.style.name).lower() == 'title':
                        continue
                    for run in p.runs:
                        run.font.size = Pt(pt_size)
                    try:
                        p.paragraph_format.line_spacing = Pt(line_spacing_pt)
                    except Exception:
                        pass


def _get_text_width_cm(doc: Document) -> float:
    try:
        sec = doc.sections[0]
        return max(1.0, float(sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm))
    except Exception:
        # 回退为典型 A4 可用宽度约 16cm
        return 16.0


def _set_table_column_widths_cm(table, widths_cm: list[float]) -> None:
    try:
        table.autofit = False
    except Exception:
        pass
    for col_idx, width in enumerate(widths_cm):
        for cell in table.columns[col_idx].cells:
            try:
                cell.width = Cm(width)
            except Exception:
                pass


def _set_table_widths_fit_doc(doc: Document, table, first_col_cm: float = 2.0) -> None:
    total_cm = _get_text_width_cm(doc)
    remainder = max(1.0, total_cm - first_col_cm)
    other_each = max(1.0, remainder / 2.0)
    _set_table_column_widths_cm(table, [first_col_cm, other_each, other_each])


def _align_table_cells(table, is_header: bool) -> None:
    # 对齐规则：表头全居中；数据行第3列右对齐，其余居中
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # 垂直居中
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass
            # 段落水平对齐
            desired = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx > 0:  # body
                if c_idx == 2:
                    desired = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    desired = WD_ALIGN_PARAGRAPH.CENTER
            for p in cell.paragraphs:
                p.alignment = desired


def insert_detail_table_at_placeholder(doc: Document, placeholder: str, headers: List[str], rows: List[List[str]], logger: logging.Logger) -> bool:
    # 优先在段落中替换占位符
    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            table = doc.add_table(rows=1, cols=len(headers))
            _fill_table(table, headers, rows)
            _apply_table_style_or_borders(table)
            # 列宽：首列窄，其余按文本宽度均分
            _set_table_widths_fit_doc(doc, table, first_col_cm=2.0)
            _center_table(table)
            _align_table_cells(table, is_header=True)
            # 将表格插入到占位段落之后
            p_elm = p._element
            tbl_elm = table._element
            p_elm.addnext(tbl_elm)
            # 删除占位段落
            p_elm.getparent().remove(p_elm)
            _apply_font_size_to_document(doc, pt_size=14, line_spacing_pt=20)
            return True

    # 若段落未命中，再在表格中查找
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        inner = cell.add_table(rows=1, cols=len(headers))
                        _fill_table(inner, headers, rows)
                        _apply_table_style_or_borders(inner)
                        _set_table_widths_fit_doc(doc, inner, first_col_cm=2.0)
                        _center_table(inner)
                        _align_table_cells(inner, is_header=True)
                        p.text = ""
                        _apply_font_size_to_document(doc, pt_size=14, line_spacing_pt=25)
                        return True
    return False


def render_docx_from_template(template_path: Path, output_path: Path, placeholders: dict, logger: logging.Logger) -> Path:
    doc = Document(str(template_path))
    replace_placeholders(doc, placeholders, logger)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("已生成文档: %s", output_path)
    return output_path


def attach_detail_table(docx_path: Path, placeholder: str, headers: List[str], rows: List[List[str]], logger: logging.Logger) -> None:
    doc = Document(str(docx_path))
    inserted = insert_detail_table_at_placeholder(doc, placeholder, headers, rows, logger)
    if not inserted:
        # 未命中占位符则追加到文末
        table = doc.add_table(rows=1, cols=len(headers))
        _fill_table(table, headers, rows)
        _apply_table_style_or_borders(table)
        _set_table_widths_fit_doc(doc, table, first_col_cm=2.0)
        _center_table(table)
        _align_table_cells(table, is_header=True)
        logger.warning("未找到占位符 %s，已将明细表追加到文档末尾", placeholder)
    # 文末追加落款日期
    from datetime import datetime
    p = doc.add_paragraph(datetime.now().strftime("%Y年%m月%d日"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _apply_font_size_to_document(doc, pt_size=14, line_spacing_pt=25)
    doc.save(str(docx_path))
    logger.info("已插入明细表: %s", docx_path)
